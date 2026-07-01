#!/usr/bin/env python3
"""
读取docx文档内容
"""
from docx import Document
import os

def read_docx_content(file_path):
    """
    读取docx文档的所有段落内容
    
    参数:
    file_path (str): docx文件路径
    
    返回:
    list: 包含所有段落文本的列表
    """
    try:
        doc = Document(file_path)
        content = []
        
        # 读取所有段落
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                content.append(para.text)
        
        # 读取所有表格
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_content.append(" | ".join(row_text))
            if table_content:
                content.extend(table_content)
        
        return content
    except Exception as e:
        print(f"读取文档时出错: {e}")
        return []

def save_to_md(content, output_path):
    """
    将内容保存为markdown文件
    
    参数:
    content (list): 要保存的内容列表
    output_path (str): 输出文件路径
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        print(f"内容已成功保存到: {output_path}")
    except Exception as e:
        print(f"保存文件时出错: {e}")

if __name__ == "__main__":
    docx_path = "d:\\Pycharm\\Pycharm_edited\\my_2nd_agent\\agent 笔记.docx"
    
    if not os.path.exists(docx_path):
        print(f"错误：文件不存在 - {docx_path}")
    else:
        print(f"正在读取文档: {docx_path}")
        content = read_docx_content(docx_path)
        
        if content:
            md_path = "d:\\Pycharm\\Pycharm_edited\\my_2nd_agent\\agent 笔记_转换版.md"
            save_to_md(content, md_path)
            
            # 打印前10行内容预览
            print("\n文档内容预览(前10行):")
            print("=" * 50)
            for i, line in enumerate(content[:10], 1):
                print(f"{i:2d}. {line}")
            if len(content) > 10:
                print(f"... 还有 {len(content) - 10} 行内容")
        else:
            print("文档内容为空或读取失败")